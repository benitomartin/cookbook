"""
document.convert_format — Convert a document between formats.

Mutable: requires user confirmation (writes a file).
Supports conversions between txt, md, html, docx.
PDF generation deferred to create_pdf tool.
"""

from __future__ import annotations

import os
from pathlib import Path

from pydantic import BaseModel, Field

from mcp_base import MCPError, MCPResult, MCPTool, ErrorCodes
from validation import assert_sandboxed, assert_absolute_path


class Params(BaseModel):
    """Parameters for convert_format."""

    path: str = Field(description="Path to source document")
    target_format: str = Field(description="Target format (txt, md, html, docx)")
    output_path: str | None = Field(default=None, description="Output path (auto-generated if omitted)")


class Result(BaseModel):
    """Return value for convert_format."""

    path: str


class ConvertFormat(MCPTool[Params, Result]):
    """Convert a document between formats."""

    name = "document.convert_format"
    description = "Convert a document between formats"
    confirmation_required = True
    undo_supported = False

    async def execute(self, params: Params) -> MCPResult[Result]:
        """Convert the source document to the target format."""
        assert_absolute_path(params.path, "path")
        assert_sandboxed(params.path)

        if not os.path.exists(params.path):
            raise MCPError(ErrorCodes.FILE_NOT_FOUND, f"File not found: {params.path}")

        allowed_formats = ("txt", "md", "html", "docx")
        target = params.target_format.lower().lstrip(".")
        if target not in allowed_formats:
            raise MCPError(
                ErrorCodes.INVALID_PARAMS,
                f"Unsupported target format: {target}. Allowed: {', '.join(allowed_formats)}",
            )

        # Determine output path
        output_path = params.output_path
        if not output_path:
            src = Path(params.path)
            output_path = str(src.with_suffix(f".{target}"))

        assert_absolute_path(output_path, "output_path")
        assert_sandboxed(output_path)

        try:
            # Extract source text
            source_ext = Path(params.path).suffix.lower()
            source_text = _read_source(params.path, source_ext)

            # Convert and write
            _write_target(source_text, output_path, target)

            return MCPResult(success=True, data=Result(path=output_path))

        except MCPError:
            raise
        except Exception as e:
            raise MCPError(ErrorCodes.INTERNAL_ERROR, f"Failed to convert: {e}") from e


def _read_source(file_path: str, ext: str) -> str:
    """Read source document as plain text."""
    if ext == ".pdf":
        from pypdf import PdfReader

        reader = PdfReader(file_path)
        return "\n\n".join(p.extract_text() or "" for p in reader.pages)

    if ext == ".docx":
        from docx import Document

        doc = Document(file_path)
        return "\n".join(p.text for p in doc.paragraphs)

    # txt, md, html, csv, etc.
    return Path(file_path).read_text(encoding="utf-8")


def _write_target(text: str, output_path: str, target_format: str) -> None:
    """Write text in the target format."""
    os.makedirs(os.path.dirname(output_path), exist_ok=True)

    if target_format == "docx":
        from docx import Document

        doc = Document()
        for paragraph in text.split("\n"):
            if paragraph.strip():
                doc.add_paragraph(paragraph)
        doc.save(output_path)
    else:
        # txt, md, html — write directly
        if target_format == "html" and not text.strip().startswith("<"):
            import markdown

            html_content = markdown.markdown(text)
            text = f"<!DOCTYPE html>\n<html>\n<body>\n{html_content}\n</body>\n</html>"

        Path(output_path).write_text(text, encoding="utf-8")
