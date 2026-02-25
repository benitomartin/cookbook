"""
document.create_docx — Create a Word document from structured content.

Mutable: requires user confirmation (writes a file).
Converts markdown-style text to a .docx using python-docx.
"""

from __future__ import annotations

import os
import re
from pathlib import Path
from typing import Any

from pydantic import BaseModel, Field

from mcp_base import MCPError, MCPResult, MCPTool, ErrorCodes
from validation import assert_sandboxed, assert_absolute_path


class Params(BaseModel):
    """Parameters for create_docx."""

    content: str = Field(description="Markdown content")
    template: str | None = Field(default=None, description="Template path")
    output_path: str = Field(description="Where to save the docx")


class Result(BaseModel):
    """Return value for create_docx."""

    path: str


class CreateDocx(MCPTool[Params, Result]):
    """Create a Word document from structured content."""

    name = "document.create_docx"
    description = "Create a Word document from structured content"
    confirmation_required = True
    undo_supported = False

    async def execute(self, params: Params) -> MCPResult[Result]:
        """Generate a DOCX from markdown content."""
        assert_absolute_path(params.output_path, "output_path")
        assert_sandboxed(params.output_path)

        try:
            from docx import Document

            doc: Document
            if params.template and os.path.exists(params.template):
                assert_sandboxed(params.template)
                doc = Document(params.template)
            else:
                doc = Document()

            # Parse markdown-style content into paragraphs
            _populate_docx(doc, params.content)

            # Ensure output directory exists
            os.makedirs(os.path.dirname(params.output_path), exist_ok=True)

            doc.save(params.output_path)

            return MCPResult(success=True, data=Result(path=params.output_path))

        except MCPError:
            raise
        except Exception as e:
            raise MCPError(ErrorCodes.INTERNAL_ERROR, f"Failed to create DOCX: {e}") from e


def _populate_docx(doc: Any, markdown_text: str) -> None:
    """Convert basic markdown to DOCX paragraphs.

    Uses duck-typing: doc must support add_heading(), add_paragraph(), save().
    """
    lines = markdown_text.split("\n")

    for line in lines:
        stripped = line.strip()

        if not stripped:
            continue

        # Headings
        heading_match = re.match(r"^(#{1,6})\s+(.+)$", stripped)
        if heading_match:
            level = len(heading_match.group(1))
            text = heading_match.group(2)
            doc.add_heading(text, level=min(level, 9))
            continue

        # Bullet lists
        bullet_match = re.match(r"^[-*]\s+(.+)$", stripped)
        if bullet_match:
            doc.add_paragraph(bullet_match.group(1), style="List Bullet")
            continue

        # Numbered lists
        num_match = re.match(r"^\d+\.\s+(.+)$", stripped)
        if num_match:
            doc.add_paragraph(num_match.group(1), style="List Number")
            continue

        # Regular paragraph
        doc.add_paragraph(stripped)
