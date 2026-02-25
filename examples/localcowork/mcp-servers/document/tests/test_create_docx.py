"""Tests for document.create_docx tool."""

from __future__ import annotations

from pathlib import Path

import pytest

from tools.create_docx import CreateDocx


@pytest.fixture()
def tool() -> CreateDocx:
    return CreateDocx()


async def test_create_docx_basic(tool: CreateDocx, tmp_dir: Path) -> None:
    """Should create a DOCX from markdown content."""
    out_path = str(tmp_dir / "output.docx")

    result = await tool.execute(
        tool.get_params_model()(
            content="# Test Document\n\nSome content here.\n\n- Bullet one\n- Bullet two",
            output_path=out_path,
        )
    )

    assert result.success is True
    assert result.data is not None
    assert result.data.path == out_path
    assert Path(out_path).exists()

    # Verify content
    from docx import Document

    doc = Document(out_path)
    texts = [p.text for p in doc.paragraphs]
    assert "Test Document" in texts
    assert "Some content here." in texts


async def test_create_docx_headings(tool: CreateDocx, tmp_dir: Path) -> None:
    """Should convert markdown headings to DOCX headings."""
    out_path = str(tmp_dir / "headings.docx")

    result = await tool.execute(
        tool.get_params_model()(
            content="# H1\n## H2\n### H3\nBody text",
            output_path=out_path,
        )
    )

    assert result.success is True

    from docx import Document

    doc = Document(out_path)
    heading_texts = [p.text for p in doc.paragraphs if p.style.name.startswith("Heading")]
    assert "H1" in heading_texts
    assert "H2" in heading_texts


async def test_create_docx_creates_dirs(tool: CreateDocx, tmp_dir: Path) -> None:
    """Should create parent directories."""
    out_path = str(tmp_dir / "nested" / "dir" / "doc.docx")

    result = await tool.execute(
        tool.get_params_model()(content="Hello", output_path=out_path)
    )

    assert result.success is True
    assert Path(out_path).exists()


def test_metadata(tool: CreateDocx) -> None:
    """Should have correct metadata."""
    assert tool.name == "document.create_docx"
    assert tool.confirmation_required is True
    assert tool.undo_supported is False
